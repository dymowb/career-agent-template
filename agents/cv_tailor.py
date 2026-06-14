"""
agents/cv_tailor.py — ATS-Optimized CV Tailoring Agent

Produces a per-role tailored CV as markdown + DOCX (Word).
DOCX is preferred by ATS parsers over PDF for text extraction.

What the agent may do:
  - Reorder bullets within each job to prioritize JD-relevant ones
  - Lightly reword bullet framing/lead-ins to align with JD language (metrics/outcomes/tech names locked)
  - Set the CV headline using your fixed identity prefix (see config / rule 7) with a tailored domain cluster
  - Write a 2-3 sentence professional summary from verified facts

What the agent may NOT do (enforced by prompt + factual judge):
  - Add any bullet not in resume_source_of_truth.md
  - Change any metric, number, or technology name
  - Invent experience, projects, or achievements
  - Remove any job or education entry
"""
import json
from pathlib import Path

import anthropic
from pydantic import BaseModel, Field

import config


# ── Output schema ─────────────────────────────────────────────────────────────

class CVExperience(BaseModel):
    company: str = ""
    title: str = ""
    period: str = ""
    bullets: list[str] = Field(default_factory=list)


class CVEducation(BaseModel):
    degree: str = ""
    institution: str = ""
    period: str = ""
    notes: str = ""


class CVProject(BaseModel):
    name: str = ""
    bullets: list[str] = Field(default_factory=list)


class CVTailorResult(BaseModel):
    headline: str = ""
    professional_summary: str = ""
    experience: list[CVExperience] = Field(default_factory=list)
    education: list[CVEducation] = Field(default_factory=list)
    certifications: list[str] = Field(default_factory=list)
    technical_projects: list[CVProject] = Field(default_factory=list)
    keywords_inserted: list[str] = Field(default_factory=list)


# ── Prompt ────────────────────────────────────────────────────────────────────

SYSTEM = """\
You are an expert ATS-optimization specialist. Produce a tailored CV for a specific role.

HARD CONSTRAINTS — a factual auditor will validate every claim:
1. NEVER add a bullet that does not exist in the source of truth (verbatim or near-verbatim)
2. NEVER change any metric, number, or technology name
3. NEVER invent experience, projects, or achievements
4. NEVER remove entire jobs or education entries
4b. NEVER reorder jobs — always output them in reverse-chronological order, exactly as
    they appear in the source of truth.
# ── OPTIONAL: candidate-specific narrative rules ───────────────────────────────
# The original author pinned a few specific bullets so the LLM could never drop or
# misorder them. Add your own here if some bullet is critical to your story, e.g.:
#   4c. The FIRST bullet of <Most Recent Job> MUST always be "<your flagship bullet>" —
#       it sets up context the other bullets depend on.
#   4d. The "<your strategic/leadership bullet>" MUST always be included unless tailoring
#       notes explicitly say to de-emphasise it.
# Delete this block if you don't need pinned bullets.
5. MAY reorder bullets within a job to put the most JD-relevant ones first,
   but keep thematically related bullets together — never interleave people
   management bullets with technical system bullets. Within each job, group in
   this order: (a) platform/systems/technical achievements, (b) AI work,
   (c) strategic/org initiatives (3YP, repositioning, executive-level programs),
   (d) people & org leadership (team growth, hiring, promotions),
   (e) cost/efficiency/infrastructure. Lead with whichever group is most
   relevant to the JD, but never scatter bullets from the same group.
6. MAY lightly reword the framing or lead-in of a bullet to better align with the JD's language,
   provided the underlying facts, all metrics, all technology names, and all outcomes remain
   identical. Only the framing/context words may change — never the substance.
   - Right: source says "Led launch of the billing API platform" and JD is about API products
     → "Built and launched the billing API platform, enabling 20+ partner teams..." (same facts,
     framing shifted to the API angle)
   - Right: source says "Led zero-downtime cloud migration" + JD uses "platform modernisation"
     → "Led zero-downtime platform modernisation of the legacy orchestrator..." (JD phrasing, same facts)
   - Wrong: adding a metric not in the source ("achieved 99.9% SLA" if source doesn't say that)
   - Wrong: changing a technology name or tool ("Kubernetes" → "container orchestration")
   - Wrong: softening or inflating an outcome ("reduced by ~40%" → "nearly halved" counts as inflation)
7. MUST set the CV headline to: "<FIXED IDENTITY PREFIX> — [domain cluster]"
   - Choose ONE identity prefix that describes you and keep it FIXED across every CV
     (e.g. "Engineering Manager", "Engineering & Business Manager", "Staff Software Engineer").
   - The domain cluster after the dash is the only tailored part. It should reflect the
     candidate's actual strengths as they relate to the JD (e.g. "Distributed Systems & Platform",
     "Payments & Infrastructure", "AI Systems & Platform Engineering").
   - The cluster should describe what the candidate brings, NOT copy the job title verbatim.
   - Keep it to 2–4 words. Do not append filler just to pad it.
8. MAY write a 2–3 sentence professional_summary using ONLY facts from the source of truth.
   Do NOT claim deep expertise or specialisation in a domain (e.g. "security expert",
   "deep expertise in security engineering") unless the candidate's primary role is in
   that domain. Describe transferable strengths instead.
   CRITICAL: NEVER write aggregate career statistics that are not verbatim in the source —
   e.g. "15+ years of experience", "teams of 50+ engineers", "managed hundreds of engineers".
   Use only specific, verifiable facts (e.g. "grew the team by N", "served X requests/day").
   If you cannot write a factually grounded summary, leave professional_summary as "".
9. MUST include technical_projects verbatim from the source of truth — do NOT modify project names, GitHub links, or bullet content

BULLET WRITING RULE:
Every bullet must tell a short story — pair an action or initiative with its outcome or
impact. A standalone data point ("Reduced X by 20%") or a standalone action ("Implemented
CI/CD pipelines") is meaningless without the other half. When the source of truth has
related cause/effect bullets, combine them. When a single bullet already contains both,
preserve that structure.

ATS RULES:
- Use the exact JD terminology — ATS does literal string matching, not semantic matching
- Do not over-insert keywords; they must read naturally
- Prioritize the most recent and most relevant bullets

Return ONLY valid JSON — no markdown fences, no explanation.

Schema:
{
  "headline": "Engineering Manager — Distributed Systems & Platform",
  "professional_summary": "2–3 sentences max, verified facts only",
  "experience": [
    {
      "company": "Current Company",
      "title": "Engineering Manager",
      "period": "2020 – Present",
      "bullets": ["reordered / keyword-enhanced bullet from source of truth"]
    }
  ],
  "education": [
    {
      "degree": "MBA",
      "institution": "Example Business School",
      "period": "2016–2018",
      "notes": "optional honors / ranking note"
    }
  ],
  "certifications": [
    "Relevant Certification — Issuer (Year)"
  ],
  "technical_projects": [
    {
      "name": "Project Name (github.com/yourusername/project)",
      "bullets": ["verbatim bullet from source of truth"]
    }
  ],
  "keywords_inserted": ["distributed systems → inserted into cloud migration bullet"]
}
"""


# ── LLM call ──────────────────────────────────────────────────────────────────

def _parse_cv_json(raw: str) -> CVTailorResult | None:
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        print("[cv_tailor] WARNING: malformed JSON")
        return None
    experience = [CVExperience(**e) if isinstance(e, dict) else e for e in data.get("experience", [])]
    education = [CVEducation(**e) if isinstance(e, dict) else e for e in data.get("education", [])]
    projects = [CVProject(**p) if isinstance(p, dict) else p for p in data.get("technical_projects", [])]
    data["experience"] = [e.model_dump() for e in experience]
    data["education"] = [e.model_dump() for e in education]
    data["technical_projects"] = [p.model_dump() for p in projects]
    return CVTailorResult(**data)


def tailor_cv(
    parsed_job: dict,
    tailoring_notes: dict,
    resume_source: str,
    candidate_profile: str,
) -> tuple[CVTailorResult, str, str]:
    """Returns (cv, user_msg, raw) so callers can use the conversation for retry."""
    client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)

    user_msg = f"""Resume Source of Truth (CRITICAL — your ONLY source of facts):
{resume_source}

Candidate Profile:
{candidate_profile}

Target Role:
{json.dumps(parsed_job, indent=2)}

Tailoring Guidance:
- Emphasize: {tailoring_notes.get('emphasize', [])}
- De-emphasize: {tailoring_notes.get('deemphasize', [])}
- Keyword alignment: {tailoring_notes.get('keyword_alignment', [])}
- Narrative angle: {tailoring_notes.get('narrative_angle', '')}

Produce a tailored CV for this role.
"""

    message = client.messages.create(
        model=config.MODEL,
        max_tokens=config.MAX_TOKENS,
        system=SYSTEM,
        messages=[{"role": "user", "content": user_msg}],
    )

    raw = message.content[0].text.strip()
    cv = _parse_cv_json(raw) or CVTailorResult()
    return cv, user_msg, raw


# ── Markdown renderer ─────────────────────────────────────────────────────────

CONTACT = config.CANDIDATE_CONTACT

def render_markdown(cv: CVTailorResult) -> str:
    lines = [f"# {config.CANDIDATE_NAME}"]
    if cv.headline:
        lines.append(f"**{cv.headline}**")
    lines += ["", CONTACT, ""]

    if cv.professional_summary:
        lines += ["---", "", "## Professional Summary", "", cv.professional_summary, ""]

    if cv.experience:
        lines += ["---", "", "## Experience", ""]
        for exp in cv.experience:
            lines += [
                f"**{exp.company}** — {exp.title}",
                f"*{exp.period}*",
                "",
            ]
            for b in exp.bullets:
                lines.append(f"- {b}")
            lines.append("")

    if cv.education:
        lines += ["---", "", "## Education", ""]
        for edu in cv.education:
            lines.append(f"**{edu.degree}** — {edu.institution}  ·  {edu.period}")
            if edu.notes:
                lines.append(f"*{edu.notes}*")
        lines.append("")

    if cv.certifications:
        lines += ["---", "", "## Certifications", ""]
        for cert in cv.certifications:
            lines.append(f"- {cert}")
        lines.append("")

    if cv.technical_projects:
        lines += ["---", "", "## Interests & Technical Practice", ""]
        for proj in cv.technical_projects:
            lines.append(f"**{proj.name}**")
            for b in proj.bullets:
                lines.append(f"- {b}")
            lines.append("")
        lines.append(f"All projects public on GitHub: {config.GITHUB_URL}")
        lines.append("")

    return "\n".join(lines)


# ── DOCX renderer ─────────────────────────────────────────────────────────────

def render_docx(cv: CVTailorResult, output_path: Path) -> bool:
    """Generate ATS-friendly DOCX. Returns True on success."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[cv_tailor] python-docx not installed — skipping DOCX")
        return False

    doc = Document()

    # ── Page margins (narrow for more content) ──────────────────────────────
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # ── Helpers ─────────────────────────────────────────────────────────────
    def set_font(run, size, bold=False, color=None):
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = RGBColor(*color)

    def para(text="", style="Normal", space_before=0, space_after=4):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def add_line(paragraph):
        """Add a horizontal rule below a paragraph."""
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def section_header(title: str):
        p = para(space_before=6, space_after=2)
        run = p.add_run(title.upper())
        set_font(run, 10, bold=True, color=(50, 50, 50))
        add_line(p)

    # ── Name & headline ─────────────────────────────────────────────────────
    p = para(space_before=0, space_after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config.CANDIDATE_NAME)
    set_font(run, 20, bold=True)

    if cv.headline:
        p = para(space_before=0, space_after=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cv.headline)
        set_font(run, 12, color=(60, 60, 60))

    p = para(space_before=0, space_after=6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CONTACT)
    set_font(run, 9, color=(80, 80, 80))

    # ── Professional summary ─────────────────────────────────────────────────
    if cv.professional_summary:
        section_header("Professional Summary")
        p = para(space_before=2, space_after=6)
        run = p.add_run(cv.professional_summary)
        set_font(run, 10)

    # ── Experience ───────────────────────────────────────────────────────────
    if cv.experience:
        section_header("Experience")
        for exp in cv.experience:
            # Company + title on one line, period on next
            p = para(space_before=4, space_after=1)
            run = p.add_run(f"{exp.company}")
            set_font(run, 10, bold=True)
            run2 = p.add_run(f"  —  {exp.title}")
            set_font(run2, 10)

            p = para(space_before=0, space_after=2)
            run = p.add_run(exp.period)
            set_font(run, 9, color=(90, 90, 90))

            for bullet in exp.bullets:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(bullet)
                set_font(run, 9)

    # ── Education ────────────────────────────────────────────────────────────
    if cv.education:
        section_header("Education")
        for edu in cv.education:
            p = para(space_before=3, space_after=1)
            run = p.add_run(f"{edu.degree}  —  {edu.institution}  ·  {edu.period}")
            set_font(run, 10, bold=True)
            if edu.notes:
                p = para(space_before=0, space_after=2)
                run = p.add_run(edu.notes)
                set_font(run, 9, color=(80, 80, 80))

    # ── Certifications ───────────────────────────────────────────────────────
    if cv.certifications:
        section_header("Certifications")
        for cert in cv.certifications:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(cert)
            set_font(run, 9)

    # ── Interests & Technical Practice ───────────────────────────────────────
    if cv.technical_projects:
        section_header("Interests & Technical Practice")
        for proj in cv.technical_projects:
            p = para(space_before=4, space_after=1)
            run = p.add_run(proj.name)
            set_font(run, 10, bold=True)
            for bullet in proj.bullets:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(bullet)
                set_font(run, 9)
        p = para(space_before=4, space_after=2)
        run = p.add_run(f"All projects public on GitHub: {config.GITHUB_URL}")
        set_font(run, 9, color=(80, 80, 80))

    doc.save(str(output_path))
    return True


# ── CLI helper ────────────────────────────────────────────────────────────────

def run_cv_tailor(parsed_job: dict, tailoring_notes: dict) -> tuple[str, Path | None]:
    """
    Returns (markdown_cv, docx_path_or_None).

    Runs factual judge internally — DOCX is only written when factual risk is
    not high (revise verdict surfaces as a warning but does not block the file).
    """
    from agents.factual_judge import run_factual_judge

    resume_source = (config.CONTEXT_DIR / "resume_source_of_truth.md").read_text()
    candidate_profile = (config.CONTEXT_DIR / "candidate_profile.md").read_text()

    print(f"[cv_tailor] tailoring CV for {parsed_job.get('title')} @ {parsed_job.get('company')}…")
    cv, user_msg, raw = tailor_cv(parsed_job, tailoring_notes, resume_source, candidate_profile)

    if not cv.experience:
        print("[cv_tailor] WARNING: LLM returned empty experience — skipping")
        return "", None

    markdown = render_markdown(cv)

    # Factual judge must approve before DOCX is written — one retry on block/revise
    factual = run_factual_judge(markdown)
    print(f"[cv_tailor] factual check: verdict={factual.verdict} risk={factual.factual_risk}")

    if (factual.verdict in ("block", "revise") or factual.factual_risk == "high") and factual.issues:
        print(f"[cv_tailor] factual issues — retrying with corrections: {factual.issues}")
        issues_text = "\n".join(f"- {i}" for i in factual.issues)
        correction_msg = (
            f"The CV you produced has the following factual issues — fix them and return corrected JSON:\n"
            f"{issues_text}\n\n"
            f"Rules: remove or correct any claim not verbatim in the source of truth. "
            f"If a professional_summary claim cannot be verified, set professional_summary to \"\"."
        )
        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
        retry_msg = client.messages.create(
            model=config.MODEL,
            max_tokens=config.MAX_TOKENS,
            system=SYSTEM,
            messages=[
                {"role": "user", "content": user_msg},
                {"role": "assistant", "content": raw},
                {"role": "user", "content": correction_msg},
            ],
        )
        retried = _parse_cv_json(retry_msg.content[0].text.strip())
        if retried and retried.experience:
            cv = retried
            markdown = render_markdown(cv)
        else:
            print("[cv_tailor] retry parse failed — using original")

        factual = run_factual_judge(markdown)
        print(f"[cv_tailor] factual re-check: verdict={factual.verdict} risk={factual.factual_risk}")
        if factual.verdict == "block" or factual.factual_risk == "high":
            print(f"[cv_tailor] BLOCKED after retry: {factual.issues}")
            return markdown, None

    if factual.verdict == "revise":
        print(f"[cv_tailor] factual concerns (non-blocking): {factual.issues}")

    role_dir = config.application_dir(
        parsed_job.get("company", "Unknown"),
        parsed_job.get("title", "Unknown"),
    )

    # Always write the markdown alongside the DOCX
    md_path = role_dir / "cv.md"
    md_path.write_text(markdown, encoding="utf-8")
    print(f"[cv_tailor] markdown → {md_path}")

    docx_path = role_dir / "cv.docx"
    if render_docx(cv, docx_path):
        print(f"[cv_tailor] DOCX → {docx_path}")
        return markdown, docx_path

    return markdown, None
